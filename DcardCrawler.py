import requests
import time
import json
import datetime
import re, urllib
import os
from bs4 import BeautifulSoup
import docx


def get_web_page(url):
    resp = requests.get(
        url=url,
        cookies={'over18': '1'}
    )
    if resp.status_code != 200:
        print('Invalid url:', resp.url)
        return None
    else:
        return resp.text

def getcontent(id):
    #print("033:"+url_content+str(id))
    post=get_web_page(url_content+str(id))
    #print("031:"+post)
    dcarddata = json.loads(post)
    return dcarddata["content"]

def save(img_urls,dname,id):
    if img_urls:
        os.makedirs(dname)
        print("005:"+dname)
        Dcarddoc.add_heading(dname)
        dcontent=getcontent(id)
        Dcarddoc.add_paragraph(dcontent)
        try:
            #print("001:"+img_urls)

            for img_url in img_urls:
                print("011:" + img_url['url'])
                dimg_url = img_url['url']
                fname = img_url['url'].split('/')[-1]
                urllib.request.urlretrieve(dimg_url, os.path.join(dname, fname))
                Dcarddoc.add_picture(fname)
        except Exception as e:
            print(e)
            print("013")




url ='https://www.dcard.tw/_api/forums/sex/posts?popular=false'
url_content='http://dcard.tw/_api/posts/'
blacklist=["[公告] 表特板板規 (2015.2.12)","[公告] 不願上表特 ＆ 優文推薦 ＆ 檢舉建議專區"]

dateoffset=1
likecount_threshod=200


def DcardTransferS2D(sDate):
    datetransfer = re.compile(r'(\d{4})-(\d{1,2})-(\d{2})')
    mo = datetransfer.search(sDate)
    if mo:
        sDcardDate = mo.group(1) + '_' + mo.group(2) + '_' + mo.group(3)  # 組合日期格式給資料夾檔名使用
        DcardDate = datetime.datetime.strptime(sDcardDate, "%Y_%m_%d").date()
    else:
        print("spttdate is abnormal:" + sDate)
    return sDcardDate,DcardDate


currentdate = datetime.date.today()
#print("002:"+str(currentdate))

Dcarddoc = docx.Document()
keeptracking=1
dcardurl=url
while keeptracking==1:
    post = get_web_page(dcardurl)
    print("CS Url:"+dcardurl)

    dcarddata = json.loads(post)
    slatestdate, latestdate = DcardTransferS2D(dcarddata[0]['updatedAt'])
    for data in dcarddata:
        if (currentdate - datetime.timedelta(days=dateoffset) <= latestdate):
            if(data['likeCount'] > likecount_threshod):
                dname = slatestdate + data['title']  # 用 strip() 去除字串前後的空白
                did=data["id"]
                print("033:" + url_content + str(did))
                #dexcerpt=data['excerpt']
                save(data['media'],dname,did)
        else:
            keeptracking=0
    if keeptracking==1:
        dcardurl=url+"&before="+str(dcarddata[29]["id"])
Dcarddoc.save('dcardata.docx')


